"""Convert non-Markdown uploads (PDF / DOCX / images) to Markdown.

Pipeline ported and adapted from ToHeinAC/MD-maker (Apache-2.0):
https://github.com/ToHeinAC/MD-maker

- PDF: per-page routing — digital-text pages are LLM-rewritten into clean
  Markdown; scanned/empty pages are rasterized and OCR'd via a vision model.
- DOCX: deterministic heading/list/table mapping (no LLM).
- Images: OCR via the vision model.

All Ollama access goes through ollama_client; all prompts live in prompts.py.
"""

import base64
import io
import os
from collections.abc import Callable, Iterator
from pathlib import Path

from dotenv import load_dotenv

import ollama_client
from prompts import OCR_DEEPSEEK_PROMPT, OCR_SYSTEM_PROMPT, OCR_USER_PROMPT, MD_REWRITE_PROMPT

load_dotenv()

OCR_MODEL = os.getenv("OCR_MODEL", "deepseek-ocr:3b")
# The per-page rewrite only adds Markdown structure and never rewords, so a small
# fixed model is ~8x faster at equal fidelity (99% word recall, no hallucination)
# and keeps PDF conversion independent of the main chat model (OLLAMA_MODEL).
REWRITE_MODEL = os.getenv("REWRITE_MODEL", "LiquidAI/lfm2.5-1.2b-instruct:latest")
PDF_DPI = int(os.getenv("PDF_DPI", "150"))

TEXT_THRESHOLD = 40  # chars; below this a PDF page is treated as image-only

CONVERTIBLE_EXTS = {
    ".pdf", ".docx",
    ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp",
}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}

# Callback signature: on_progress(done: int, total: int, label: str) -> None
ProgressCb = Callable[[int, int, str], None]


def is_convertible(filename: str) -> bool:
    """True if the file extension can be converted to Markdown."""
    return Path(filename).suffix.lower() in CONVERTIBLE_EXTS


# --- image / PDF helpers (from MD-maker pdf_utils.py) ----------------------

def _image_to_base64(pil_image) -> str:
    buf = io.BytesIO()
    pil_image.convert("RGB").save(buf, format="JPEG", quality=85)
    return base64.b64encode(buf.getvalue()).decode()


def iter_pdf_pages(pdf_bytes: bytes, dpi: int = PDF_DPI) -> Iterator[tuple[str, object]]:
    """Yield ('text', str) for pages with extractable text, else ('image', PIL.Image)."""
    import pypdfium2 as pdfium

    scale = dpi / 72
    doc = pdfium.PdfDocument(pdf_bytes)
    for page in doc:
        textpage = page.get_textpage()
        text = textpage.get_text_range().strip()
        if len(text) >= TEXT_THRESHOLD:
            yield "text", text
        else:
            bitmap = page.render(scale=scale)
            yield "image", bitmap.to_pil()


def _pdf_page_count(pdf_bytes: bytes) -> int:
    import pypdfium2 as pdfium

    return len(pdfium.PdfDocument(pdf_bytes))


# --- LLM calls (from MD-maker models.py) -----------------------------------

def rewrite_text(text: str, model_id: str = REWRITE_MODEL) -> str:
    """Reformat already-extracted PDF text into Markdown without altering wording."""
    return ollama_client.rewrite(model_id, MD_REWRITE_PROMPT + text)


def convert_image(pil_image, model_id: str = OCR_MODEL) -> str:
    """OCR one image to Markdown via a vision model."""
    img_b64 = _image_to_base64(pil_image)
    if model_id.startswith("deepseek-ocr"):
        prompt = OCR_DEEPSEEK_PROMPT
    else:
        prompt = f"{OCR_SYSTEM_PROMPT}\n\n{OCR_USER_PROMPT}"
    return ollama_client.ocr(model_id, prompt, img_b64)


# --- DOCX helpers (from MD-maker docx_utils.py) ----------------------------

def _paragraph_md(para) -> str | None:
    text = para.text.strip()
    if not text:
        return None
    style = para.style.name
    if style.startswith("Heading 1"):
        return f"# {text}"
    if style.startswith("Heading 2"):
        return f"## {text}"
    if style.startswith("Heading 3"):
        return f"### {text}"
    if "Bullet" in style:
        return f"- {text}"
    if "Number" in style:
        return f"1. {text}"
    return text


def _table_md(table) -> str | None:
    rows = table.rows
    if not rows:
        return None
    lines = []
    header = [c.text.strip() for c in rows[0].cells]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def extract_docx_text(docx_bytes: bytes) -> str:
    """Convert a .docx file to Markdown, preserving document order (no LLM step)."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            md = _paragraph_md(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            md = _table_md(Table(child, doc))
        else:
            md = None
        if md:
            parts.append(md)
    return "\n\n".join(parts)


# --- unified entry point ---------------------------------------------------

def _convert_pdf(pdf_bytes: bytes, on_progress: ProgressCb | None) -> str:
    total = _pdf_page_count(pdf_bytes)
    parts: list[str] = []
    for i, (kind, payload) in enumerate(iter_pdf_pages(pdf_bytes)):
        if on_progress:
            on_progress(i, total, f"Page {i + 1}/{total} ({'OCR' if kind == 'image' else 'text'})")
        if kind == "text":
            parts.append(rewrite_text(payload))
        else:
            parts.append(convert_image(payload))
    if on_progress:
        on_progress(total, total, "Done")
    return "\n\n".join(parts)


def convert_to_markdown(
    file_bytes: bytes, filename: str, on_progress: ProgressCb | None = None
) -> str:
    """Convert a supported file to a Markdown string. Dispatch by extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        try:
            return _convert_pdf(file_bytes, on_progress)
        finally:
            # Free the OCR model so the ingest step (a text model) gets the full
            # GPU instead of sharing/splitting VRAM with the resident vision model.
            ollama_client.unload(OCR_MODEL)
    if suffix == ".docx":
        if on_progress:
            on_progress(0, 1, "Converting DOCX")
        md = extract_docx_text(file_bytes)
        if on_progress:
            on_progress(1, 1, "Done")
        return md
    if suffix in IMAGE_EXTS:
        from PIL import Image

        if on_progress:
            on_progress(0, 1, "OCR image")
        try:
            md = convert_image(Image.open(io.BytesIO(file_bytes)))
        finally:
            ollama_client.unload(OCR_MODEL)  # free VRAM for the ingest text model
        if on_progress:
            on_progress(1, 1, "Done")
        return md
    raise ValueError(f"Unsupported file type: {suffix}")
