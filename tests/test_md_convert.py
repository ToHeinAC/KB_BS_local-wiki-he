"""Tests for md_convert.py — non-Markdown → Markdown conversion."""

import io

import pytest

import md_convert


# --- is_convertible --------------------------------------------------------

@pytest.mark.parametrize("name", ["a.pdf", "A.PDF", "b.docx", "c.png", "d.JPG", "e.tiff"])
def test_is_convertible_true(name):
    assert md_convert.is_convertible(name) is True


@pytest.mark.parametrize("name", ["a.md", "b.txt", "c.html", "d", "e.csv"])
def test_is_convertible_false(name):
    assert md_convert.is_convertible(name) is False


# --- DOCX (deterministic, no LLM) ------------------------------------------

def _make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_heading("Section", level=2)
    doc.add_paragraph("Body text.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_headings_and_paragraph():
    md = md_convert.extract_docx_text(_make_docx())
    assert "# Title" in md
    assert "## Section" in md
    assert "Body text." in md


def test_extract_docx_table():
    md = md_convert.extract_docx_text(_make_docx())
    assert "| A | B |" in md
    assert "| --- | --- |" in md
    assert "| 1 | 2 |" in md


def test_convert_to_markdown_docx_routes_to_extractor():
    md = md_convert.convert_to_markdown(_make_docx(), "doc.docx")
    assert "# Title" in md


# --- PDF routing (monkeypatched, no real model / no pypdfium2) --------------

def test_pdf_routing_text_and_image(monkeypatch):
    pages = [("text", "raw text page"), ("image", object())]
    monkeypatch.setattr(md_convert, "_pdf_page_count", lambda b: len(pages))
    monkeypatch.setattr(md_convert, "iter_pdf_pages", lambda b, dpi=md_convert.PDF_DPI: iter(pages))
    monkeypatch.setattr(md_convert, "rewrite_text", lambda t: f"REWRITE[{t}]")
    monkeypatch.setattr(md_convert, "convert_image", lambda img: "OCR")

    md = md_convert.convert_to_markdown(b"fake-pdf", "x.pdf")
    assert md == "REWRITE[raw text page]\n\nOCR"


def test_pdf_progress_callback(monkeypatch):
    pages = [("text", "a"), ("text", "b")]
    monkeypatch.setattr(md_convert, "_pdf_page_count", lambda b: len(pages))
    monkeypatch.setattr(md_convert, "iter_pdf_pages", lambda b, dpi=md_convert.PDF_DPI: iter(pages))
    monkeypatch.setattr(md_convert, "rewrite_text", lambda t: t)
    calls = []
    md_convert.convert_to_markdown(b"x", "x.pdf", on_progress=lambda d, t, l: calls.append((d, t)))
    assert calls[-1] == (2, 2)  # final "Done" tick


# --- image OCR routing ------------------------------------------------------

def test_image_routes_to_convert_image(monkeypatch):
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (2, 2), "white").save(buf, format="PNG")
    monkeypatch.setattr(md_convert, "convert_image", lambda img: "IMG-OCR")
    md = md_convert.convert_to_markdown(buf.getvalue(), "scan.png")
    assert md == "IMG-OCR"


# --- LLM wrappers select the right prompt ----------------------------------

def test_convert_image_uses_deepseek_prompt(monkeypatch):
    seen = {}
    monkeypatch.setattr(md_convert, "_image_to_base64", lambda img: "b64")
    monkeypatch.setattr(
        md_convert.ollama_client, "ocr",
        lambda model, prompt, img_b64: seen.update(model=model, prompt=prompt) or "ok",
    )
    md_convert.convert_image(object(), model_id="deepseek-ocr:3b")
    assert seen["prompt"] == md_convert.OCR_DEEPSEEK_PROMPT


def test_convert_image_uses_system_prompt_for_non_deepseek(monkeypatch):
    seen = {}
    monkeypatch.setattr(md_convert, "_image_to_base64", lambda img: "b64")
    monkeypatch.setattr(
        md_convert.ollama_client, "ocr",
        lambda model, prompt, img_b64: seen.update(prompt=prompt) or "ok",
    )
    md_convert.convert_image(object(), model_id="some-vision:1b")
    assert md_convert.OCR_SYSTEM_PROMPT in seen["prompt"]


def test_unsupported_extension_raises():
    with pytest.raises(ValueError):
        md_convert.convert_to_markdown(b"x", "file.csv")
